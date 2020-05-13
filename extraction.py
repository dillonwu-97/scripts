######################################################################
# Extraction program that reads in a cover letter and fill its in
# Converts the cover letter to both docx and pdf format
# Usage:
# cat info.txt | python extraction.py read_in.docx output.docx
######################################################################

import docx
from docx import Document
import argparse
from collections import Counter
import string
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
import os



def getText (filename):

	#Extracting all the paragraphs in the document
	fullText = []
	for p in filename.paragraphs:
		fullText.append(p.text)

	#key, value pair where key is thing to replace, value is what to replace it with
	fill_In = Counter()
	for para in fullText:
		for x in para.split():
			if "$" in x:
				fill_In[x.strip(".,")] = ""
				#fill_In.
	
	#putting the stuff into the dictionary
	for i in fill_In:
		temp = input("Please enter the: " + i)
		fill_In[i] = temp
	
	print(fill_In)

	#running it back one more time to edit the paragraphs
	for para in range(len(fullText)):
		#print(type(para))
		for word in fill_In:
			fullText[para] = fullText[para].replace(word, fill_In[word])
			#print(para)
	return fullText

def main(): 

	parser = argparse.ArgumentParser()
	parser.add_argument("doc", help = "file to be read in")
	parser.add_argument("output", help = "output name")
	args = parser.parse_args()
	filename = docx.Document(args.doc)
	fullText = getText(filename)
	
	#Creating the document
	testdoc = Document()

	#Storing page 1
	section = testdoc.sections[0]

	#Formatting the header
	header = section.header
	p = header.paragraphs[0]
	p.text = 'Dillon Wu \n Emory University - Clairmont MSC 161517; 1762 Clifton Road; Atlanta, GA 30322-1006 \n dillon.wu@emory.edu • (1917) 826 8232'
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	section.left_margin = Inches(1)
	section.right_margin = Inches(1)

	#Need to do carriage return 
	temp = ""
	for i in range (0, 3):
		if (i != 2):
			temp = temp + fullText[i] + '\r'
		else: temp = temp + fullText[i]
	testdoc.add_paragraph(temp)

	for i in range(3, len(fullText)):
		j = testdoc.add_paragraph(fullText[i])
	testdoc.save('/Users/Kvothe/Desktop/work_applications_19:20/'+ args.output)
	os.system('cd /Users/Kvothe/Desktop/work_applications_19:20/; docx2pdf ' + args.output)


if __name__ == "__main__":
	main()

